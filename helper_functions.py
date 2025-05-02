import os
from dotenv import load_dotenv
from pydantic import Field, BaseModel
from typing import List, Optional, Literal
from io import BytesIO
from helper import get_resume_details, get_resume_details_from_image
from PyPDF2 import PdfReader
from docx import Document
import docx2txt

from langchain_core.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser, PydanticOutputParser
from helper import IsResume, SkillGaps

def load_env_variables():

    load_dotenv()

    os.environ["OPENAI_API_KEY"]=os.getenv("OPENAI_API_KEY")
    os.environ["TAVILY_API_KEY"]=os.getenv("TAVILY_API_KEY")

    gemini_api_key=os.getenv("GEMINI_API_KEY")
    youtube_api_key=os.getenv("YOUTUBE_API_KEY")
    groq_api_key=os.getenv("GROQ_API_KEY")

    return gemini_api_key, youtube_api_key, groq_api_key

def load_file(uploaded_file):
    if uploaded_file is not None:
        file_extension = os.path.splitext(uploaded_file)[1].lower()
        file_buffer = uploaded_file

        if file_extension == ".pdf":
            reader = PdfReader(file_buffer)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            
        elif file_extension in [".docx", ".doc"]:
            try:
                # Try multiple methods to extract text from Word docs
                text = ""
                try:
                    # Method 1: python-docx
                    doc = Document(file_buffer)
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    
                    # Get text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    paragraphs.append(cell.text.strip())
                    
                    text = "\n".join(paragraphs)
                except:
                    print("python-docx failed, trying docx2txt...")
                
                # If text is still empty, try docx2txt
                if not text.strip():
                    text = docx2txt.process(file_buffer)
                
                if not text.strip():
                    raise ValueError("Could not extract text from document")
                    
                return text.strip()
                
            except Exception as e:
                print(f"Error processing Word document: {e}")
                return None
                
        elif file_extension in [".png", ".jpg", ".jpeg"]:
            text=get_resume_details_from_image(image_file_path=file_buffer)
            
        return text
    
    return None
    
def get_suggestions(text,agent,api_keys: dict):
    try:
        # Get career summary from resume
        career_summary = get_resume_details(text=text, model='gemini')
        if not career_summary:
            raise ValueError("Could not extract career details from resume")

        # Initialize models
        model = ChatGoogleGenerativeAI(
            model='gemini-2.0-flash',
            api_key=api_keys["GEMINI_API_KEY"]
        )

        # Extract job roles
        job_roles = career_summary.job_role
        if not job_roles:
            raise ValueError("No job roles found in resume")

        # Get industry trends
        industry_trends = agent.invoke({
            "messages": f"Do a detailed search for the required skillsets and current trend for the job roles {job_roles} and provide your answers in a clean format."
        })
        
        if not industry_trends or "messages" not in industry_trends:
            raise ValueError("Could not fetch industry trends")

        industry_trends_results = str(industry_trends.get("messages")[-2].content) + str(industry_trends.get("messages")[-1].content)

        # Analyze skill gaps
        parser_skill_gaps = PydanticOutputParser(pydantic_object=SkillGaps)
        agent_prompt_template = PromptTemplate(
            template="""
                You will be given a career summary of the candidate which is as follows: \n{career_summary}.\n 
                You will also be given required skillsets and industry trends for one or more than one job roles {job_roles} 
                which is as follows: {industry_trends_results}. \n 
                You need to compare the candidates profile with the industry trends and find out the skill gaps, 
                strengths, weaknesses, areas of improvement.\n {format_instructions}
            """,
            input_variables=["career_summary", "job_roles", "industry_trends_results"],
            partial_variables={
                "format_instructions": parser_skill_gaps.get_format_instructions()
            }
        )

        skill_gaps_chain = agent_prompt_template | model | parser_skill_gaps
        skill_gaps = skill_gaps_chain.invoke({
            "career_summary": career_summary,
            "job_roles": job_roles,
            "industry_trends_results": industry_trends_results
        })

        # Summarize skill gaps
        str_parser = StrOutputParser()
        summarize_skill_gaps_prompt = PromptTemplate(
            template="""
                You will receive details about a candidate's weaknesses: {weaknesses} 
                and areas of improvement: {areas_of_improvement}. 
                You will also get the job role/roles that suit(s) the candidate: {job_roles}. 
                You need to summarize the candidate's weaknesses and areas of improvement in a single paragraph.
            """,
            input_variables=["weaknesses", "areas_of_improvement", "job_roles"]
        )

        model_groq = ChatGroq(
            model="qwen-qwq-32b",
            api_key=api_keys["GROQ_API_KEY"]
        )

        summarize_skill_gaps_chain = summarize_skill_gaps_prompt | model_groq | str_parser
        skill_gaps_summary = summarize_skill_gaps_chain.invoke({
            "weaknesses": skill_gaps.weaknesses,
            "areas_of_improvement": skill_gaps.areas_of_improvement,
            "job_roles": job_roles
        })

        # Fetch learning resources
        agent_prompt_template_fetch_materials = """
            You will receive summarized details about a candidate's weaknesses and areas of improvement: {skill_gaps_summary}. 
            You will also get the job role/roles that suit(s) the candidate: {job_roles}. 
            You need to fetch relevant courses or resources or information according to the candidate's weaknesses 
            and areas of improvement and the job role. Use the tools at your disposal. Use multiple tools if and when required. 
            Provide courses, materials from the internet and youtube video links and also provide URLs for each. 
            Use all the tools given to you: `youtube_search_tool`, `text_generator_tool`, `duckduckgo_search_tool`, 
            `coursera_search_tool` and `wiki_tool`. Address the candidate as a second person.
        """

        resources = agent.invoke({
            "messages": agent_prompt_template_fetch_materials.format(
                skill_gaps_summary=skill_gaps_summary, 
                job_roles=job_roles
            )
        })

        return resources

    except Exception as e:
        print(f"Error in get_suggestions: {str(e)}")
        raise